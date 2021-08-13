import win32com.client as win32
import os
import re
from win32com.client import constants
import pandas as pd
from docx import Document
from openpyxl import Workbook
from openpyxl import load_workbook
import numpy as np


def doc_to_docx(dir):
    for r, d, f in os.walk(dir):
        for file in f:
            if ".doc" in file:
                convert_to_docx(os.path.join(os.path.realpath(r), file))


def convert_to_docx(path):
    word = win32.gencache.EnsureDispatch('Word.Application')
    document = word.Documents.Open(path)
    document.Activate()

    new_file = os.path.abspath(path)
    new_file = re.sub(r'\.\w+$', '.docx', new_file)

    word.ActiveDocument.SaveAs(
        new_file,
        FileFormat=constants.wdFormatXMLDocument
    )
    document.Close(False)
    os.remove(path)


def main(dir):
    current_dir = os.getcwd()

    for r, d, f in os.walk(dir):
        for file in f:
            if ".docx" in file:
                doc = Document(os.path.join(r, file).split('/', 1)[1])
                city = get_city_name(doc)
                full_path = os.path.join(os.path.realpath(r), file)
                rename_dir(current_dir, city, full_path)
                doc = Document(city+'.docx')
                data = read_table(doc)
                write_excel(data, city)


def get_city_name(doc):
    if len(doc.paragraphs[0].text) > 26:
        # Get city if the paragraph is working
        city_name = get_city_by_paragraph(doc)
    else:
        city_name = get_city_by_xml(doc)

    return format_city_name(city_name)


def get_city_by_paragraph(doc):
    # Get text from first paragraph of the document
    paragraph = doc.paragraphs[0].text
    # Split and format text
    city = paragraph.split(' ', 4)[4].lower()

    return city


def get_city_by_xml(doc):
    # Get all document content
    data = doc._element.xpath('.//w:t')
    # Get city name
    city = data[1].text
    # Split and format text
    city = city.split(' ', 1)[1].lower()

    return city


def format_city_name(city):
    prepositions = ['de', 'do']
    items = []
    for item in city.split():
        if item not in prepositions:
            item = item.capitalize()
        items.append(item)
    return ' '.join(items)


def rename_dir(current_dir, city, old_name):
    new_name = current_dir + '\\' + city + '.docx'
    if os.path.isfile(new_name):
        print("The file already exists")
    else:
        os.rename(old_name, new_name)


def read_table(document, table=9):
    # Choose table in document
    table = document.tables[table-1]
    # Put data from table in a array
    data = [[cell.text for cell in row.cells] for row in table.rows]
    # Delete first line and select just the columns 2 and 5
    data = np.delete(data, 0, axis=0)
    data = np.array(data[:, [2, 5]])
    # Delete items 7 and 10 | flatten (axis=None) the array
    data = np.delete(data, [7, 10], axis=None)
    # Swap positions 5 and 6
    data[5], data[6] = data[6], data[5]
    for i in range(14):
        data[i] = data[i].replace(',', '.')
        if not is_numeric(data[i]):
            data[i] = '0'

    data = data.astype(float)
    print(data)
    return data


def is_numeric(s):
    try:
        float(s)
        return True
    except ValueError:
        return False


def write_excel(data, city):
    wb = load_workbook('Preços de  03 a 09  de agosto de 2021 (1).xlsm', keep_vba=True)
    ws = wb['Preços']
    print(city)
    for j in range(8, 28):
        if ws.cell(row=j, column=1).value == city:
            for i in range(1, 15):
                if data[i-1] != 0:
                    ws.cell(row=j, column=i+1).value = data[i-1]
    wb.save('Preços de  03 a 09  de agosto de 2021 (1).xlsm')


# Start point
directory = './'
doc_to_docx(directory)
main(directory)
